import streamlit as st
import io, re, smtplib, ssl, html
import base64
from email.message import EmailMessage
from pathlib import Path
from docx import Document
from pdfminer.high_level import extract_text
from email_validator import validate_email, EmailNotValidError

# Copia por defecto a la cátedra si no hay secreto TEACHER_EMAIL / TEACHER_BCC.
DEFAULT_TEACHER_EMAIL = "investigacion@uccuyo.edu.ar"

# ---------------------------------
# Configuración
# ---------------------------------
st.set_page_config(page_title="Auto-corrección | Metodología", layout="centered")

RUBRIC_MAX = {1: 100, 2: 100, 3: 100, 4: 100, 5: 100, 6: 100, 7: 100, 8: 100}

PRACTICO_LABELS = {
    1: "Práctico Nº 1 — IA en la escritura del proyecto",
    2: "Práctico Nº 2 — Establecimiento de Métodos de Recolección de Datos y Tipos de Muestreos. Tamaño de la Muestra",
    3: "Práctico Nº 3 — Operacionalización de Variables y Determinación de Métodos de Análisis de Datos",
    4: "Práctico Nº 4 — Introducción + Marco teórico + Búsqueda (≈500 palabras en total)",
    5: "Trabajo práctico Módulo 5 — Mendeley: citas en Word y bibliografía",
    6: "Trabajo práctico Módulo 6 — Estilos de Word e índice automático",
    7: "Práctico Nº 7 — Análisis cuantitativo",
    8: "Práctico Nº 8 — Análisis cualitativo",
}

# ---------------------------------
# Lectura de archivos
# ---------------------------------
def read_docx(file_bytes: bytes) -> dict:
    bio = io.BytesIO(file_bytes)
    doc = Document(bio)
    paragraphs, texts = [], []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        style = getattr(p.style, "name", "") or ""
        if txt:
            paragraphs.append((txt, style))
            texts.append(txt)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                paragraphs.append((line, "Table"))
                texts.append(line)
    return {"plain_text": "\n".join(texts), "paragraphs": paragraphs, "filetype": "docx"}

def read_pdf(file_bytes: bytes) -> dict:
    bio = io.BytesIO(file_bytes)
    text = extract_text(bio) or ""
    return {
        "plain_text": text,
        "paragraphs": [(line.strip(), "") for line in text.splitlines() if line.strip()],
        "filetype": "pdf",
    }

def parse_file(uploaded) -> dict:
    suffix = Path(uploaded.name).suffix.lower()
    file_bytes = uploaded.read()
    if suffix == ".docx":
        return read_docx(file_bytes)
    elif suffix == ".pdf":
        return read_pdf(file_bytes)
    else:
        st.error("Formato no soportado. Suba un archivo .docx o .pdf")
        return {"plain_text": "", "paragraphs": [], "filetype": "unknown"}

# ---------------------------------
# Utilidades de evaluación
# ---------------------------------
def count_in_text(patterns, text_lower):
    return sum(1 for p in patterns if p in text_lower)

def apa_inline_citations(text):
    """Cuenta citas APA frecuentes: (Apellido, 2020), (A y B, 2020), (A et al., 2020), Apellido (2020)."""
    patterns = [
        r"\([A-Za-zÁÉÍÓÚÜÑáéíóúüñ\-]+(?:\s+y\s+[A-Za-zÁÉÍÓÚÜÑáéíóúüñ\-]+)?,\s*(?:19|20)\d{2}\)",
        r"\([A-Za-zÁÉÍÓÚÜÑáéíóúüñ\-]+\s+et\s+al\.,\s*(?:19|20)\d{2}\)",
        r"\b[A-ZÁÉÍÓÚÜÑ][a-záéíóúüñ\-]+\s+\((?:19|20)\d{2}\)",
    ]
    found = set()
    for pat in patterns:
        for m in re.finditer(pat, text):
            found.add(m.group(0).lower())
    return len(found)

def mentions_ai(text_lower):
    """Detecta mención real de IA, no la sílaba 'ia' dentro de 'social' o 'ciencia'."""
    t = (text_lower or "").lower()
    frases = [
        "inteligencia artificial",
        "chatgpt",
        "chat gpt",
        "herramienta de ia",
        "herramientas de ia",
        "uso de ia",
        "apoyo de ia",
        "ia generativa",
        "copilot",
        "gemini",
        "claude",
    ]
    if any(f in t for f in frases):
        return True
    return re.search(r"(^|[^a-záéíóúüñ])ia([^a-záéíóúüñ]|$)", t) is not None

def has_bibliography_section(text_lower):
    keys = ["bibliografía", "referencias", "referencias bibliográficas"]
    return any(k in text_lower for k in keys)

def find_headings_docx(paragraphs):
    """Cuenta títulos por estilo en DOCX."""
    h1 = sum(1 for _, s in paragraphs if "Heading 1" in s or "Título 1" in s)
    h2 = sum(1 for _, s in paragraphs if "Heading 2" in s or "Título 2" in s)
    h3 = sum(1 for _, s in paragraphs if "Heading 3" in s or "Título 3" in s)
    return h1, h2, h3

def has_toc(text_lower, paragraphs, filetype):
    """Detecta índice automático; no alcanza con la palabra suelta 'contenido'."""
    keys = [
        "tabla de contenido",
        "tabla de contenidos",
        "índice automático",
        "indice automatico",
        "table of contents",
    ]
    if any(k in text_lower for k in keys):
        return True
    if filetype == "docx":
        if any("table of contents" in p[0].lower() for p in paragraphs):
            return True
        if any("toc" in (s or "").lower() for _, s in paragraphs):
            return True
    return False

def build_feedback_message(num, score, breakdown, summary):
    lines = []
    lines.append("Resultado de la corrección automática:\n")
    lines.append(f"{PRACTICO_LABELS[num]}")
    lines.append(f"Puntaje: {score}/{RUBRIC_MAX[num]}\n")
    lines.append("Desglose por criterios:")
    for name, got, mx, expl in breakdown:
        lines.append(f" - {name}: {got}/{mx}. {expl}")
    lines.append("\nComentarios generales:")
    lines.append(summary if summary else "—")
    return "\n".join(lines)

# ---------------------------------
# Rúbricas por práctico
# ---------------------------------
def corregir_practico_1(text, paragraphs, filetype):
    t = text.lower(); total = 0; bd = []
    # Tema y Título (20)
    found = count_in_text(["tema", "título"], t)
    if found >= 2: pts, expl = 20, "Se identificaron 'Tema' y 'Título'."
    elif found == 1: pts, expl = 10, "Solo se encontró uno (Tema o Título)."
    else: pts, expl = 0, "No se detectaron secciones claras de 'Tema' y 'Título'."
    total += pts; bd.append(("Tema y Título", pts, 20, expl))
    # Paradigma (15)
    pts = 15 if "paradigma" in t else 0
    total += pts; bd.append(("Paradigma", pts, 15, "Incluye el paradigma de investigación." if pts else "No se encontró el apartado de paradigma."))
    # Pregunta (20)
    pts = 20 if ("pregunta de investigación" in t or re.search(r"pregunta(s)?\s+de\s+investigación", t)) else 0
    total += pts; bd.append(("Pregunta de investigación", pts, 20, "Incluye pregunta de investigación." if pts else "No se detectó una pregunta explícita."))
    # Objetivos (30)
    has_general = "objetivo general" in t or "objetivo principal" in t
    has_especificos = "objetivos específicos" in t or "objetivos especificos" in t
    if has_general and has_especificos: pts, expl = 30, "Incluye objetivo general/principal y objetivos específicos."
    elif has_general or has_especificos: pts, expl = 15, "Solo se encontró uno (general/principal o específicos)."
    else: pts, expl = 0, "No se detectaron objetivos claros."
    total += pts; bd.append(("Objetivos", pts, 30, expl))
    # Hipótesis (15)
    if "hipótesis" in t or "hipotesis" in t: pts, expl = 15, "Incluye hipótesis de investigación."
    else: pts, expl = 10, "Sin hipótesis explícita; se asume diseño que no la requiere."
    total += pts; bd.append(("Hipótesis (si corresponde)", pts, 15, expl))
    return total, bd, "Se evaluó la presencia de secciones fundamentales de un anteproyecto."

def corregir_practico_2(text, paragraphs, filetype):
    t = text.lower(); total = 0; bd = []
    # Operacionalización (45)
    keys = ["variable","independiente","dependiente","definición conceptual","definicion conceptual",
            "definición operacional","definicion operacional","indicador","escala","instrumento",
            "unidad de análisis","unidades de análisis"]
    found = count_in_text(keys, t)
    if found >= 7: pts, expl = 45, "Se identifican los campos centrales del cuadro."
    elif found >= 4: pts, expl = 30, "Cuadro parcialmente completo."
    else: pts, expl = 10, "No se reconoce un cuadro completo."
    total += pts; bd.append(("Cuadro de operacionalización", pts, 45, expl))
    # Métodos (25)
    methods = ["análisis","regresión","correlación","anova","t-student","chi-cuadrado",
               "temático","codificación","grounded theory","análisis de contenido","estadístico","cualitativo"]
    hits = count_in_text(methods, t)
    if hits >= 3: pts, expl = 25, "Describe métodos de análisis y su pertinencia."
    elif hits >= 1: pts, expl = 15, "Menciona métodos con poco detalle."
    else: pts, expl = 5, "No especifica métodos de análisis."
    total += pts; bd.append(("Métodos de análisis", pts, 25, expl))
    # Validación (20)
    val = ["validez","fiabilidad","confiabilidad","triangulación","alfa de cronbach","pilotaje","validación de instrumentos"]
    hits = count_in_text(val, t)
    if hits >= 2: pts, expl = 20, "Estrategias de validación para datos/instrumentos."
    elif hits == 1: pts, expl = 10, "Menciona validación de forma breve."
    else: pts, expl = 0, "No indica cómo validará datos/instrumentos."
    total += pts; bd.append(("Validación de datos/instrumentos", pts, 20, expl))
    # Ética (10)
    pts = 10 if any(k in t for k in ["ética","consentimiento informado","anonimato","confidencialidad"]) else 0
    total += pts; bd.append(("Ética", pts, 10, "Incluye consideraciones éticas." if pts else "No se describen consideraciones éticas."))
    return total, bd, "Se verificó cuadro de variables, pertinencia de métodos, validación y ética."

def corregir_practico_3(text, paragraphs, filetype):
    t = text.lower(); total = 0; bd = []
    # Muestreo (30)
    ms = ["muestreo","probabilístico","no probabilístico","aleatorio","estratificado",
          "intencionado","conglomerados","bola de nieve","sistemático"]
    hits = count_in_text(ms, t)
    if hits >= 2 and ("fundament" in t or "justific" in t):
        pts, expl = 30, "Describe el muestreo y fundamenta la elección."
    elif hits >= 1:
        pts, expl = 20, "Menciona el tipo de muestreo con poca justificación."
    else:
        pts, expl = 10, "No se identifica claramente el muestreo."
    total += pts; bd.append(("Tipo de muestreo y justificación", pts, 30, expl))
    # Instrumentos (25)
    inst = ["cuestionario","encuesta","entrevista","guía","observación","escala","test"]
    hits = count_in_text(inst, t)
    if hits >= 2: pts, expl = 25, "Selecciona instrumentos y explica su adecuación."
    elif hits == 1: pts, expl = 15, "Menciona un instrumento sin suficiente justificación."
    else: pts, expl = 5, "No define instrumentos de recolección."
    total += pts; bd.append(("Instrumentos y adecuación", pts, 25, expl))
    # Validez/fiabilidad (20)
    val = ["validez","fiabilidad","confiabilidad","pilotaje","alfa de cronbach"]
    hits = count_in_text(val, t)
    if hits >= 2: pts, expl = 20, "Incluye procedimientos para validez/fiabilidad."
    elif hits == 1: pts, expl = 10, "Menciona brevemente validez/fiabilidad."
    else: pts, expl = 0, "No aborda validez/fiabilidad."
    total += pts; bd.append(("Validez/fiabilidad de instrumentos", pts, 20, expl))
    # Tamaño muestral (25)
    tm = ["tamaño de la muestra","n=","muestra de","cálculo muestral","error","confianza"]
    hits = count_in_text(tm, t)
    if hits >= 2: pts, expl = 25, "Estima tamaño de muestra y fundamenta (error/confianza/supuestos)."
    elif hits == 1: pts, expl = 15, "Menciona el tamaño sin fundamento claro."
    else: pts, expl = 5, "No calcula ni fundamenta el tamaño muestral."
    total += pts; bd.append(("Tamaño de la muestra", pts, 25, expl))
    return total, bd, "Se revisaron muestreo, instrumentos, validez y tamaño muestral."

def corregir_practico_4(text, paragraphs, filetype):
    """Se espera ~500 palabras EN TOTAL para Introducción+Marco (±10%). ≥3 citas, mención de IA y bibliografía."""
    t = text.lower(); total = 0; bd = []
    intro_idx = t.find("introducción"); marco_idx = t.find("marco teórico")
    total_words = len(re.findall(r"\w+", text))
    # Extensión total
    if intro_idx != -1 and marco_idx != -1:
        pts = 20 if 450 <= total_words <= 550 else (10 if 350 <= total_words <= 650 else 0)
        expl = f"Extensión total {total_words} palabras (objetivo ~500)."
    else:
        pts = 10 if 450 <= total_words <= 550 else 0
        expl = "No se detectaron ambos títulos; se evaluó por extensión total."
    total += pts; bd.append(("Extensión total (Intro+Marco)", pts, 20, expl))
    # Citas
    citas = apa_inline_citations(text)
    if citas >= 3: pts, expl = 30, f"Se detectaron {citas} citas (mínimo 3)."
    elif citas == 2: pts, expl = 20, "Solo se detectaron 2 citas."
    elif citas == 1: pts, expl = 10, "Solo se detectó 1 cita."
    else: pts, expl = 0, "No se detectaron citas (Apellido, Año)."
    total += pts; bd.append(("Citas en el texto", pts, 30, expl))
    # IA (15)
    pts = 15 if mentions_ai(t) else 0
    total += pts; bd.append(("Uso de IA (mención)", pts, 15, "Se menciona uso de IA." if pts else "No se menciona explícitamente apoyo de IA."))
    # Bibliografía (15)
    pts = 15 if has_bibliography_section(t) else 0
    total += pts; bd.append(("Referencias/Bibliografía", pts, 15, "Incluye bibliografía." if pts else "No se detectó sección de bibliografía."))
    # Búsqueda de fuentes (20) — el práctico lo pide y antes no sumaba al 100
    busq = [
        "búsqueda", "busqueda", "scielo", "scopus", "scholar", "pubmed",
        "redalyc", "dialnet", "base de datos", "bases de datos",
        "palabras clave", "descriptores", "repositorio", "fuentes",
    ]
    hits = count_in_text(busq, t)
    if hits >= 2: pts, expl = 20, "Describe la búsqueda de fuentes o bases consultadas."
    elif hits == 1: pts, expl = 10, "Menciona la búsqueda de forma breve."
    else: pts, expl = 0, "No se reconoce el apartado de búsqueda de fuentes."
    total += pts; bd.append(("Búsqueda de fuentes", pts, 20, expl))
    return total, bd, "Se evaluó extensión total (~500), citas, mención de IA, bibliografía y búsqueda."

def corregir_practico_5(text, paragraphs, filetype):
    t = text.lower(); total = 0; bd = []
    citas = apa_inline_citations(text)
    if citas >= 5: pts, expl = 35, f"Se detectaron {citas} citas (mínimo 5)."
    elif citas >= 3: pts, expl = 20, f"Solo {citas} citas; se requieren 5."
    elif citas >= 1: pts, expl = 10, "Muy pocas citas."
    else: pts, expl = 0, "No se detectaron citas."
    total += pts; bd.append(("Citas en el texto", pts, 35, expl))
    pts = 30 if has_bibliography_section(t) else 10
    total += pts; bd.append(("Bibliografía final", pts, 30, "Incluye bibliografía generada." if pts == 30 else "No se detecta bibliografía clara."))
    has_year = len(re.findall(r"(19|20)\d{2}", text)) >= 5
    has_doi_or_url = "doi" in t or "http" in t
    if has_year and (has_doi_or_url or "vol." in t or "pp." in t or "nº" in t or "no." in t):
        pts, expl = 20, "Referencias con metadatos y formato consistente."
    else:
        pts, expl = 10, "Formato poco consistente o incompleto."
    total += pts; bd.append(("Consistencia de formato", pts, 20, expl))
    pts = 15 if any(k in t for k in ["mendeley","carpeta","grupo","metadatos","corrigiendo metadatos"]) else 5
    total += pts; bd.append(("Organización/metadatos", pts, 15, "Evidencia organización/corrección de metadatos." if pts == 15 else "No se menciona organización/metadatos."))
    return total, bd, "Se verificaron citas, bibliografía final y consistencia general de referencias."

def corregir_practico_6(text, paragraphs, filetype):
    t = text.lower(); total = 0; bd = []
    if filetype == "docx":
        h1,h2,h3 = find_headings_docx(paragraphs)
        if h1>=1 and h2>=1 and h3>=1: pts, expl = 50, f"Jerarquía correcta: H1={h1}, H2={h2}, H3={h3}."
        elif (h1>=1 and h2>=1) or (h1>=1 and h3>=1): pts, expl = 35, f"Faltan algunos niveles: H1={h1}, H2={h2}, H3={h3}."
        else: pts, expl = 15, f"Escasa jerarquía: H1={h1}, H2={h2}, H3={h3}."
    else:
        caps = len(re.findall(r"\n[A-ZÁÉÍÓÚÑ ]{6,}\n", "\n"+text+"\n"))
        pts, expl = (35, "Jerarquías aproximadas en PDF.") if caps>=3 else ((20, "Jerarquía mínima en PDF.") if caps>=1 else (10, "No se reconoce jerarquía en PDF."))
    total += pts; bd.append(("Títulos jerarquizados", pts, 50, expl))
    toc = has_toc(t, paragraphs, filetype)
    pts, expl = (40, "Se detecta tabla de contenido/índice.") if toc else (15, "No se detecta índice automático.")
    total += pts; bd.append(("Tabla de contenido", pts, 40, expl))
    pts = 10 if any(k in t for k in ["actualizar índice","actualizar el índice","update table of contents"]) else 5
    total += pts; bd.append(("Actualización del índice (mención)", pts, 10, "Menciona la actualización del índice." if pts==10 else "No se menciona actualización."))
    return total, bd, "Se evaluó estructura por niveles y presencia de índice automático."

def corregir_practico_7(text, paragraphs, filetype):
    """Cuantitativo: descriptivas, p-value/Mann-Whitney, Cronbach, Spearman o clúster (25 c/u)."""
    t = text.lower(); total = 0; bd = []
    desc_hits = sum(1 for k in ["media","mediana","moda","desvío estándar","desvio estandar"] if k in t)
    if desc_hits>=4: pts, expl = 25, "Incluye media, mediana, moda y desvío estándar."
    elif desc_hits>=2: pts, expl = 15, "Incluye parte de las medidas descriptivas."
    elif desc_hits>=1: pts, expl = 8, "Menciona al menos una medida."
    else: pts, expl = 0, "No se reconocen medidas descriptivas."
    total += pts; bd.append(("Medidas descriptivas", pts, 25, expl))
    has_p = bool(re.search(r"p\s*[<=>]\s*0\.\d+", t))
    has_mw = ("mann-whitney" in t) or ("mann whitney" in t) or ("u de mann" in t)
    if has_p and has_mw: pts, expl = 25, "Reporta p-value y menciona Mann-Whitney."
    elif has_p or has_mw: pts, expl = 18, "Menciona prueba de significancia (p-value o Mann-Whitney)."
    else: pts, expl = 6, "No se evidencia prueba de significancia."
    total += pts; bd.append(("Significancia / Mann-Whitney", pts, 25, expl))
    has_cron = ("cronbach" in t) or ("α" in text) or ("alfa" in t) or ("alpha" in t)
    cron_value = bool(re.search(r"(cronbach|α|alfa|alpha)\s*[=:]\s*0\.\d+", t))
    if cron_value: pts, expl = 25, "Informa alfa de Cronbach con valor."
    elif has_cron: pts, expl = 15, "Menciona alfa de Cronbach sin valor."
    else: pts, expl = 5, "No se reconoce evaluación de confiabilidad."
    total += pts; bd.append(("Confiabilidad (Cronbach)", pts, 25, expl))
    has_spear = ("spearman" in t) or ("ρ" in text) or ("rho" in t)
    spear_value = bool(re.search(r"(spearman|ρ|rho).{0,12}[=]\s*[-+]?\d*\.?\d+", t))
    has_cluster = ("clúster" in t) or ("cluster" in t) or ("k-means" in t) or ("agrupamiento" in t)
    if spear_value or has_cluster: pts, expl = 25, "Presenta correlación de Spearman (con valor) o análisis de clúster."
    elif has_spear: pts, expl = 15, "Menciona Spearman sin valor."
    else: pts, expl = 8, "No se reconoce correlación ni clúster."
    total += pts; bd.append(("Relación entre variables", pts, 25, expl))
    return total, bd, "Se verificaron descriptivas, significancia/Mann-Whitney, Cronbach y relación (Spearman/clúster)."

def corregir_practico_8(text, paragraphs, filetype):
    """Cualitativo: temático (30), sentimiento (25), discurso (25), interpretación MMH vs Medicina Social (20)."""
    t = text.lower(); total = 0; bd = []
    has_tema = ("análisis temático" in t) or ("analisis tematico" in t) or ("temas" in t and "subtemas" in t)
    if has_tema and ("subtema" in t or "subtemas" in t): pts, expl = 30, "Identifica temas y subtemas con evidencia."
    elif has_tema: pts, expl = 20, "Menciona análisis temático de forma general."
    else: pts, expl = 8, "No se reconoce análisis temático."
    total += pts; bd.append(("Análisis temático", pts, 30, expl))
    sent_hits = sum(k in t for k in ["sentimiento","positiv","negativ","neutral"])
    if sent_hits>=3: pts, expl = 25, "Clasifica sentimientos con ejemplos."
    elif sent_hits>=1: pts, expl = 15, "Menciona sentimiento de forma parcial."
    else: pts, expl = 6, "No se reconoce análisis de sentimiento."
    total += pts; bd.append(("Análisis de sentimiento", pts, 25, expl))
    disco_hits = 0
    for k in ["análisis del discurso","analisis del discurso","función descriptiva","funcion descriptiva",
              "función explicativa","funcion explicativa","poder","posicionamiento","ideolog","complejid"]:
        if k in t: disco_hits += 1
    if disco_hits>=3: pts, expl = 25, "Analiza funciones del discurso, complejidad y relaciones de poder/posicionamientos."
    elif disco_hits>=1: pts, expl = 15, "Aborda el discurso parcialmente."
    else: pts, expl = 6, "No se reconoce análisis del discurso."
    total += pts; bd.append(("Análisis del discurso", pts, 25, expl))
    has_mmh = "modelo médico hegemónico" in t or "modelo medico hegemonico" in t
    has_ms = "medicina social" in t
    if has_mmh and has_ms: pts, expl = 20, "Interpreta resultados vinculándolos con MMH y Medicina Social."
    elif has_mmh or has_ms: pts, expl = 12, "Refiere a MMH o Medicina Social, sin contraste claro."
    else: pts, expl = 6, "No vincula la interpretación con MMH/Medicina Social."
    total += pts; bd.append(("Interpretación (MMH vs Medicina Social)", pts, 20, expl))
    return total, bd, "Se revisó temático, sentimiento, discurso e interpretación MMH vs Medicina Social."

# ---------------------------------
# Router de evaluación
# ---------------------------------
def evaluar_practico(num, text, paragraphs, filetype):
    if num == 1: return corregir_practico_1(text, paragraphs, filetype)
    # Intercambio de lógica: 
    # - Práctico 2 debe evaluar muestreo, instrumentos, validación y tamaño muestral
    # - Práctico 3 debe evaluar operacionalización y métodos de análisis
    if num == 2: return corregir_practico_3(text, paragraphs, filetype)
    if num == 3: return corregir_practico_2(text, paragraphs, filetype)
    if num == 4: return corregir_practico_4(text, paragraphs, filetype)
    if num == 5: return corregir_practico_5(text, paragraphs, filetype)
    if num == 6: return corregir_practico_6(text, paragraphs, filetype)
    if num == 7: return corregir_practico_7(text, paragraphs, filetype)
    if num == 8: return corregir_practico_8(text, paragraphs, filetype)
    return 0, [], "—"

# ---------------------------------
# Envío de correo (SendGrid + fallback Gmail SMTP)
# ---------------------------------
def _secret(name, default=None):
    try:
        if name in st.secrets:
            val = st.secrets[name]
            if val is None:
                return default
            val = str(val).strip()
            return val or default
    except Exception:
        pass
    return default

def correos_catedra():
    """Correos de la cátedra: TEACHER_EMAIL, CATEDRA_EMAIL y/o TEACHER_BCC (coma-separados)."""
    found = []
    for key in ("TEACHER_EMAIL", "CATEDRA_EMAIL", "TEACHER_BCC"):
        val = _secret(key)
        if val:
            found.extend(x.strip() for x in val.split(",") if x.strip())
    if not found:
        found.append(DEFAULT_TEACHER_EMAIL)
    uniq, seen = [], set()
    for email in found:
        key = email.lower()
        if key not in seen:
            seen.add(key)
            uniq.append(email)
    return uniq

def enviar_email(destinatario, asunto, mensaje_texto):
    """
    Envía la devolución AL ALUMNO (To) y una COPIA A LA CÁTEDRA (Cc).
    1) SendGrid  2) Gmail SMTP si SendGrid no está o falla.
    """
    sender_email = _secret("SENDER_EMAIL") or _secret("EMAIL_USER")
    sender_name  = _secret("SENDER_NAME", "Cátedra Metodología")
    reply_to     = _secret("REPLY_TO", sender_email)
    copias = [c for c in correos_catedra() if c.lower() != (destinatario or "").lower()]

    html_body = (
        "<html><body>"
        f"<pre style='white-space:pre-wrap;font-family:inherit'>{html.escape(mensaje_texto)}</pre>"
        "<hr>"
        "<p style='font-size:12px;color:#666'>"
        "Este correo fue enviado automáticamente por la plataforma de la cátedra "
        "(copia a la cátedra). "
        "Si no lo ves en tu bandeja de entrada, revisá <b>Correo no deseado</b> y marcá como ‘No es spam’."
        "</p>"
        "</body></html>"
    )

    # Ruta 1: SendGrid
    api_key = _secret("SENDGRID_API_KEY")
    if api_key:
        try:
            from sendgrid import SendGridAPIClient
            from sendgrid.helpers.mail import Mail, Email, To, Cc, ReplyTo

            mail = Mail(
                from_email=Email(sender_email, sender_name),
                to_emails=[To(destinatario)],
                subject=asunto,
                html_content=html_body,
                plain_text_content=mensaje_texto,
            )
            for copia in copias:
                mail.add_cc(Cc(copia))
            if reply_to:
                mail.reply_to = ReplyTo(reply_to)

            sg = SendGridAPIClient(api_key)
            resp = sg.send(mail)
            if resp.status_code in (200, 202):
                return True
            else:
                st.warning(f"SendGrid devolvió status {resp.status_code}; pruebo SMTP Gmail…")
        except Exception as e:
            st.warning(f"Fallo SendGrid: {e}; pruebo SMTP Gmail…")

    # Ruta 2: Gmail SMTP (fallback)
    try:
        remitente = _secret("EMAIL_USER")
        password  = _secret("EMAIL_PASS")
        if not remitente or not password:
            raise RuntimeError("Faltan EMAIL_USER / EMAIL_PASS en secretos de Streamlit.")

        msg = EmailMessage()
        msg["From"] = f"{sender_name} <{sender_email or remitente}>"
        msg["To"] = destinatario
        msg["Subject"] = asunto
        if reply_to:
            msg["Reply-To"] = reply_to
        if copias:
            msg["Cc"] = ", ".join(copias)
        msg.set_content(mensaje_texto)
        msg.add_alternative(html_body, subtype="html")

        smtp_server = "smtp.gmail.com"
        port = 587  # STARTTLS
        context = ssl.create_default_context()
        with smtplib.SMTP(smtp_server, port) as server:
            server.ehlo()
            server.starttls(context=context)
            server.login(remitente, password)
            server.send_message(msg)
        return True

    except Exception as e:
        st.error(f"No se pudo enviar el correo (SendGrid/SMTP). Error: {e}")
        return False

# ---------------------------------
# Interfaz Streamlit
# ---------------------------------
st.markdown(
    """
    <style>
    :root {
        --bg: #eef1f5;
        --card: #ffffff;
        --dark: #171a22;
        --text: #1f2937;
        --muted: #6b7280;
        --green-1: #0d5b53;
        --green-2: #0a7a64;
        --green-soft: #e6f4ea;
        --border: #dbe3ec;
    }
    .stApp {
        background: radial-gradient(circle at top left, #f7f9fb 0%, var(--bg) 50%, #e5eaf0 100%);
    }
    .block-container {
        max-width: 850px;
        padding-top: 2.2rem;
        padding-bottom: 2rem;
    }
    .hero {
        border-radius: 14px;
        padding: 1.25rem 1.5rem;
        background: linear-gradient(110deg, var(--green-1), var(--green-2));
        color: #fff;
        box-shadow: 0 12px 28px rgba(15, 76, 66, 0.22);
        margin-bottom: 0.9rem;
    }
    .hero-row {
        display: flex;
        align-items: center;
        gap: 1rem;
    }
    .hero-logo {
        width: 90px;
        height: 90px;
        border-radius: 10px;
        object-fit: cover;
        background: rgba(255, 255, 255, 0.16);
        padding: 4px;
    }
    .hero h1 {
        margin: 0;
        font-size: 2rem;
        line-height: 1.15;
        font-weight: 800;
    }
    .hero p {
        margin: 0.28rem 0 0;
        opacity: 0.95;
        font-size: 1rem;
    }
    .intro-card {
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 1.1rem 1.2rem;
        box-shadow: 0 8px 20px rgba(0, 0, 0, 0.05);
        margin-bottom: 0.9rem;
    }
    .intro-card h2 {
        margin: 0 0 0.45rem 0;
        font-size: 1.65rem;
        color: #111827;
    }
    .intro-card p {
        margin: 0.2rem 0;
        color: var(--muted);
    }
    .stTextInput label, .stSelectbox label, .stFileUploader label {
        font-weight: 700 !important;
        color: #111827 !important;
    }
    .stTextInput input, .stSelectbox div[data-baseweb="select"] > div {
        border-radius: 10px !important;
        border: 1px solid #c5ced8 !important;
        background: #ffffff !important;
        color: #111827 !important;
        min-height: 44px;
    }
    .stFileUploader [data-testid="stFileUploaderDropzone"] {
        border-radius: 10px;
        border: 1px dashed #9ca3af;
        background: #f8fafc;
    }
    .stButton > button {
        border: 0;
        border-radius: 10px;
        background: linear-gradient(110deg, #0f766e, #14b8a6);
        color: #fff;
        font-weight: 800;
        padding: 0.62rem 1.1rem;
        box-shadow: 0 8px 18px rgba(15, 118, 110, 0.24);
    }
    .stButton > button:hover {
        filter: brightness(1.03);
        transform: translateY(-1px);
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    f"""
    <section class="hero">
        <div class="hero-row">
            <img class="hero-logo" src="data:image/png;base64,{base64.b64encode(Path("assets/logo_uccuyo.png").read_bytes()).decode("utf-8")}" alt="Logo Universidad Católica de Cuyo" />
            <div>
                <h1>Universidad Católica de Cuyo</h1>
                <p>Secretaría de Investigación</p>
                <p>Consejo de Investigación</p>
            </div>
        </div>
    </section>
    <section class="intro-card">
        <h2>Sistema de evaluación de prácticos del Curso Metodología de la Investigación</h2>
        <p>Complete solo los campos que correspondan.</p>
        <p>El sistema genera un puntaje automático con desglose por criterios y envía la devolución al correo de quien envía el práctico y una copia a la cátedra.</p>
    </section>
    """,
    unsafe_allow_html=True,
)

correo = st.text_input("Correo electrónico del alumno")

# Select con nombres nominales
opciones = [PRACTICO_LABELS[k] for k in PRACTICO_LABELS]
label_seleccionado = st.selectbox("Práctico", opciones)
label_to_num = {v: k for k, v in PRACTICO_LABELS.items()}
practico_num = label_to_num[label_seleccionado]

uploaded = st.file_uploader("Subir archivo (.docx o .pdf)", type=["docx", "pdf"])

if st.button("Corregir y Enviar resultado"):
    if not uploaded or not correo:
        st.warning("Debe subir un archivo y un correo válido.")
    else:
        try:
            validate_email(correo)
            parsed = parse_file(uploaded)
            text, paragraphs, filetype = parsed["plain_text"], parsed["paragraphs"], parsed["filetype"]

            score, breakdown, summary = evaluar_practico(practico_num, text, paragraphs, filetype)
            mensaje = build_feedback_message(practico_num, score, breakdown, summary)

            st.metric("Puntaje", f"{score}/{RUBRIC_MAX[practico_num]}")
            st.text_area("Devolución:", mensaje, height=300)

            asunto = f"Resultado — {PRACTICO_LABELS[practico_num]}"
            if enviar_email(correo, asunto, mensaje):
                st.success(
                    f"Corregido. Devolución enviada a {correo} y copia a la cátedra."
                )
            else:
                st.warning(
                    "La corrección está arriba, pero el correo no se pudo enviar. "
                    "Revisá los secretos de Streamlit (SendGrid o Gmail) y TEACHER_EMAIL."
                )
        except EmailNotValidError:
            st.error("Correo electrónico inválido.")
